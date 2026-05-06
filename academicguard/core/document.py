"""Document parsing: plain text, DOCX, PDF, and LaTeX sources."""

from __future__ import annotations

import re
import textwrap
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class DocumentSection:
    title: str
    content: str
    level: int = 1


@dataclass
class Document:
    """Unified representation of an academic document from any source format."""

    raw_text: str
    title: str = ""
    abstract: str = ""
    keywords: list[str] = field(default_factory=list)
    sections: list[DocumentSection] = field(default_factory=list)
    references: list[str] = field(default_factory=list)
    source_path: Optional[Path] = None
    source_format: str = "text"

    # ------------------------------------------------------------------ #
    # Factory methods
    # ------------------------------------------------------------------ #

    @classmethod
    def from_file(cls, path: str | Path) -> "Document":
        p = Path(path)
        suffix = p.suffix.lower()
        if suffix == ".txt":
            return cls._from_text(p)
        if suffix == ".docx":
            return cls._from_docx(p)
        if suffix == ".pdf":
            return cls._from_pdf(p)
        if suffix in {".tex", ".latex"}:
            return cls._from_latex(p)
        raise ValueError(f"Unsupported file format: {suffix}")

    @classmethod
    def from_string(cls, text: str, title: str = "") -> "Document":
        doc = cls(raw_text=text, title=title, source_format="text")
        doc._parse_structure()
        return doc

    # ------------------------------------------------------------------ #
    # Internal parsers
    # ------------------------------------------------------------------ #

    @classmethod
    def _from_text(cls, path: Path) -> "Document":
        text = path.read_text(encoding="utf-8", errors="replace")
        doc = cls(raw_text=text, source_path=path, source_format="text")
        doc._parse_structure()
        return doc

    @classmethod
    def _from_docx(cls, path: Path) -> "Document":
        try:
            import docx
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
        d = docx.Document(str(path))
        paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        doc = cls(raw_text=text, source_path=path, source_format="docx")
        doc._parse_docx_structure(d)
        return doc

    @classmethod
    def _from_pdf(cls, path: Path) -> "Document":
        try:
            from pdfminer.high_level import extract_text as pdf_extract
        except ImportError:
            raise ImportError("Install pdfminer.six: pip install pdfminer.six")
        text = pdf_extract(str(path))
        doc = cls(raw_text=text or "", source_path=path, source_format="pdf")
        doc._parse_structure()
        return doc

    @classmethod
    def _from_latex(cls, path: Path) -> "Document":
        text = path.read_text(encoding="utf-8", errors="replace")
        clean = _strip_latex_commands(text)
        doc = cls(raw_text=clean, source_path=path, source_format="latex")
        doc._parse_latex_structure(text)
        return doc

    # ------------------------------------------------------------------ #
    # Structure extraction
    # ------------------------------------------------------------------ #

    def _parse_structure(self) -> None:
        """Heuristic section detection for plain text."""
        lines = self.raw_text.splitlines()
        abstract_lines: list[str] = []
        in_abstract = False
        in_refs = False
        current_section: list[str] = []
        current_title = ""

        for line in lines:
            stripped = line.strip()
            upper = stripped.upper()

            # Title heuristic: first non-empty short line
            if not self.title and stripped and len(stripped) < 200:
                self.title = stripped
                continue

            # Abstract
            if re.match(r"^abstract[:\s\-]*$", stripped, re.I):
                in_abstract = True
                continue
            if in_abstract:
                if re.match(r"^(keywords?|index terms?)[:\s]", stripped, re.I):
                    self.abstract = " ".join(abstract_lines).strip()
                    in_abstract = False
                    kw_text = re.sub(r"^(keywords?|index terms?)[:\s\-]*", "", stripped, flags=re.I)
                    self.keywords = [k.strip() for k in re.split(r"[;,]", kw_text) if k.strip()]
                else:
                    abstract_lines.append(stripped)
                continue

            # References section
            if re.match(r"^references?\s*$", stripped, re.I):
                if current_title and current_section:
                    self.sections.append(DocumentSection(current_title, "\n".join(current_section)))
                in_refs = True
                current_title = "References"
                current_section = []
                continue
            if in_refs:
                if stripped:
                    self.references.append(stripped)
                continue

            # Section headings: ALL CAPS or Roman numeral pattern
            is_heading = (
                re.match(r"^(I{1,3}V?|V?I{0,3}|IX|X{1,3})\.\s+\S", stripped)
                or (stripped.isupper() and 5 < len(stripped) < 80)
                or re.match(r"^\d+\.\s+[A-Z]", stripped)
            )
            if is_heading:
                if current_title and current_section:
                    self.sections.append(DocumentSection(current_title, "\n".join(current_section)))
                current_title = stripped
                current_section = []
            else:
                current_section.append(stripped)

        if current_title and current_section:
            self.sections.append(DocumentSection(current_title, "\n".join(current_section)))

    def _parse_docx_structure(self, d) -> None:
        import docx
        abstract_lines: list[str] = []
        in_abstract = False
        refs: list[str] = []
        current_title = ""
        current_body: list[str] = []

        for para in d.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            is_heading = "heading" in style_name or para.style.name.startswith("Heading")

            if not self.title and is_heading:
                self.title = text
                continue

            if re.match(r"^abstract$", text, re.I):
                in_abstract = True
                continue
            if in_abstract:
                if re.match(r"^(keywords?|index terms?)", text, re.I):
                    self.abstract = " ".join(abstract_lines)
                    in_abstract = False
                    kw = re.sub(r"^(keywords?|index terms?)[:\s\-]*", "", text, flags=re.I)
                    self.keywords = [k.strip() for k in re.split(r"[;,]", kw) if k.strip()]
                else:
                    abstract_lines.append(text)
                continue

            if re.match(r"^references?$", text, re.I):
                if current_title:
                    self.sections.append(DocumentSection(current_title, "\n".join(current_body)))
                current_title = "References"
                current_body = []
                continue

            if is_heading:
                if current_title:
                    self.sections.append(DocumentSection(current_title, "\n".join(current_body)))
                current_title = text
                current_body = []
            else:
                current_body.append(text)

        if current_title:
            self.sections.append(DocumentSection(current_title, "\n".join(current_body)))

    def _parse_latex_structure(self, latex_src: str) -> None:
        # Extract title
        m = re.search(r"\\title\{([^}]+)\}", latex_src)
        if m:
            self.title = _strip_latex_commands(m.group(1)).strip()

        # Extract abstract
        m = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", latex_src, re.S)
        if m:
            self.abstract = _strip_latex_commands(m.group(1)).strip()

        # Extract keywords
        m = re.search(r"\\keywords?\{([^}]+)\}", latex_src)
        if not m:
            m = re.search(r"\\begin\{keywords?\}(.*?)\\end\{keywords?\}", latex_src, re.S)
        if m:
            kw_raw = _strip_latex_commands(m.group(1))
            self.keywords = [k.strip() for k in re.split(r"[;,\\\\]", kw_raw) if k.strip()]

        # Extract sections
        for match in re.finditer(r"\\(sub)*section\*?\{([^}]+)\}(.*?)(?=\\(sub)*section|\Z)", latex_src, re.S):
            title = _strip_latex_commands(match.group(2)).strip()
            body = _strip_latex_commands(match.group(3)).strip()
            level = 1 + match.group(0).count("sub")
            if title.lower() in {"references", "bibliography"}:
                bib = re.findall(r"\\bibitem[^}]*\}(.*?)(?=\\bibitem|\Z)", match.group(3), re.S)
                self.references = [_strip_latex_commands(b).strip() for b in bib]
            else:
                self.sections.append(DocumentSection(title, body, level))

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #

    @property
    def body_text(self) -> str:
        """Full body text excluding abstract and references."""
        return "\n\n".join(s.content for s in self.sections if s.title.lower() != "references")

    @property
    def full_text(self) -> str:
        return self.raw_text

    @property
    def word_count(self) -> int:
        return len(self.raw_text.split())

    @property
    def sentence_count(self) -> int:
        return len(re.findall(r"[.!?]+", self.raw_text))

    def section_by_name(self, name: str) -> Optional[DocumentSection]:
        low = name.lower()
        for s in self.sections:
            if low in s.title.lower():
                return s
        return None

    def __repr__(self) -> str:
        return f"Document(title={self.title!r}, words={self.word_count}, sections={len(self.sections)})"


def _strip_latex_commands(text: str) -> str:
    """Remove common LaTeX commands, leaving plain text."""
    text = re.sub(r"\\(emph|textbf|textit|textrm|textsc|text|mathrm|cite|ref|label|url|href)\{([^}]*)\}", r"\2", text)
    text = re.sub(r"\\[a-zA-Z]+\*?\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?", " ", text)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"%[^\n]*", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()
